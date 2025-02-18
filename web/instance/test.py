from flask import Flask, render_template, request, redirect, url_for, send_file, session
import pandas as pd
import os
from datetime import datetime
from docx import Document
import zipfile
import io
app = Flask(__name__)
app.secret_key = os.urandom(24)
FILE_PATH = "data.xlsx"
TEMPLATE_PATH = "template.docx"
TEMPLATE_DIR = "templetes(1)"

def load_data():
    if not os.path.exists(FILE_PATH):
        df = pd.DataFrame(columns=['ID', 'First Name', 'Last Name', 'Fathers Name', 'Birth Date',
                                   'Age', 'Depart', 'Position', 'Rate', 'RTCK', 'Main/Part-time'])
        df.to_excel(FILE_PATH, index=False)
    return pd.read_excel(FILE_PATH)

def save_data(df):
    df.to_excel(FILE_PATH, index=False)

def calculate_age(birth_date):
    today = datetime.today().date()
    return today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/enter_data', methods=['GET', 'POST'])
def enter_data():
    message = ""
    df = load_data()
    if request.method == 'POST':
        id_input = request.form['id']
        first_name = request.form['first_name']
        last_name = request.form['last_name']
        fathers_name = request.form['fathers_name']
        birth_date = datetime.strptime(request.form['birth_date'], '%Y-%m-%d').date()
        age = calculate_age(birth_date)
        depart = request.form['depart']
        position = request.form['position']
        rate = float(request.form['rate'])
        rtck = request.form['rtck']
        main_parttime = request.form['main_parttime']
        if id_input in df['ID'].astype(str).values:
            existing_row = df[df['ID'].astype(str) == id_input].iloc[0]
            message = f"Дані з таким ID {id_input} вже існують: {existing_row['First Name']} {existing_row['Last Name']}, {existing_row['Position']}"
        else:
            new_row = pd.DataFrame([{ 'ID': id_input, 'First Name': first_name, 'Last Name': last_name, 'Fathers Name': fathers_name,
                'Birth Date': birth_date, 'Age': age, 'Depart': depart, 'Position': position,
                'Rate': rate, 'RTCK': rtck, 'Main/Part-time': main_parttime }])
            df = pd.concat([df, new_row], ignore_index=True)
            save_data(df)
            message = "Дані успішно додано!"
    return render_template('enter_data.html', message=message)


def fill_certificate(template_path, data, output_filename):
    """Заповнює шаблон документа переданими даними та зберігає його в пам'яті."""
    doc = Document(template_path)

    # Заміна тексту у всіх параграфах
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, data)

    # Заміна тексту у таблицях
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraph(cell.paragraphs[0], data)

    # Збереження документа в пам'яті
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)  # Потрібно повернутись на початок потоку
    return doc_stream

def replace_text_in_paragraph(paragraph, data):
    """Замінює всі маркери у переданому параграфі."""
    for key, value in data.items():
        placeholder = f"{{{key}}}"
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, str(value))

@app.route('/filter_data', methods=['GET', 'POST'])
def filter_data():
    df = load_data()
    filtered_df = df.copy()
    message = ""

    if request.method == 'POST':
        keywords = request.form.get('keywords', '').strip().lower()
        age_input = request.form.get('age_range', '').strip()

        if keywords:
            filtered_df = filtered_df[filtered_df.apply(lambda row: any(keywords in str(val).lower() for val in row), axis=1)]

        if age_input:
            try:
                if '-' in age_input:
                    min_age, max_age = map(int, age_input.split('-'))
                    filtered_df = filtered_df[(filtered_df['Age'] >= min_age) & (filtered_df['Age'] <= max_age)]
                else:
                    age_value = int(age_input)
                    filtered_df = filtered_df[filtered_df['Age'] == age_value]
            except ValueError:
                message = "Неправильний формат вікового введення!"

        session['filtered_data'] = filtered_df.to_dict(orient='records')

    templates = get_templates()
    return render_template('filter_data.html', kits=filtered_df.to_dict(orient='records'), message=message, templates=templates)

def get_templates():
    """Отримує список доступних шаблонів у папці."""
    if not os.path.exists(TEMPLATE_DIR):
        os.makedirs(TEMPLATE_DIR)
    return [f for f in os.listdir(TEMPLATE_DIR) if f.endswith(".docx")]

@app.route('/generate_docs', methods=['POST'])
def generate_docs():
    filtered_df = pd.DataFrame(session.get('filtered_data', []))

    if filtered_df.empty:
        return "Немає даних для генерації документа.", 400

    selected_template = request.form.get("template")
    if not selected_template:
        return "Будь ласка, виберіть шаблон.", 400

    template_path = os.path.join(TEMPLATE_DIR, selected_template)

    # Створення архіву в пам'яті
    zip_stream = io.BytesIO()

    # Створення архіву
    with zipfile.ZipFile(zip_stream, 'w', zipfile.ZIP_DEFLATED) as zipf:
        # Генерація документів і додавання їх до архіву
        for _, row in filtered_df.iterrows():
            person_data = {
                "FirstName": row["First Name"],
                "LastName": row["Last Name"],
                "FathersName": row["Fathers Name"],
                "BirthDate": row["Birth Date"],
                "Age": row["Age"],
                "Depart": row["Depart"],
                "Position": row["Position"],
                "Rate": row["Rate"],
                "RTCK": row["RTCK"],
                "MainPartTime": row["Main/Part-time"],
            }
            filename = f"{row['Last Name']}_{row['First Name']}_certificate.docx"
            doc_stream = fill_certificate(template_path, person_data, filename)
            
            # Додаємо документ до архіву
            zipf.writestr(filename, doc_stream.read())  # Додаємо файл в архів

    zip_stream.seek(0)  # Повертаємось на початок потоку для скачування

    # Повертаємо архів як вкладення
    return send_file(zip_stream, as_attachment=True, download_name="Завантаження.zip", mimetype="application/zip")


if __name__ == '__main__':
    app.run(debug=True)



#####################################
from flask import Flask, render_template, request, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from flask import request, render_template
from sqlalchemy import or_, and_

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db' 
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)


class Kit(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    first_name = db.Column(db.String(100), nullable=False)
    last_name = db.Column(db.String(100), nullable=False)
    fathers_name = db.Column(db.String(100))
    birth_date = db.Column(db.Date)
    age = db.Column(db.Integer)  # Додано стовпець для віку
    depart = db.Column(db.String(100), nullable=False)
    position = db.Column(db.String(100), nullable=False)
    rate = db.Column(db.Numeric(10, 2))
    rtck = db.Column(db.String(100))
    main_parttime = db.Column(db.String(100))


def calculate_age(birth_date):
    today = datetime.today().date()
    age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
    return age


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/enter_data', methods=['GET', 'POST'])
def enter_data():
    message = ""
    if request.method == 'POST':
        id_input = request.form['id']
        first_name = request.form['first_name']
        last_name = request.form['last_name']
        fathers_name = request.form['fathers_name']
        birth_date = datetime.strptime(request.form['birth_date'], '%Y-%m-%d').date()
        age = calculate_age(birth_date) 
        depart = request.form['depart']
        position = request.form['position']
        rate = float(request.form['rate'])
        rtck = request.form['rtck']
        main_parttime = request.form['main_parttime']

        existing_kit = Kit.query.filter_by(id=id_input).first()
        if existing_kit:
            message = f"Дані з таким id {id_input} вже існують: {existing_kit.first_name} {existing_kit.last_name}, {existing_kit.position}"
        else:
            new_kit = Kit(id=id_input, first_name=first_name, last_name=last_name, fathers_name=fathers_name,
                          birth_date=birth_date, age=age, depart=depart, position=position, rate=rate,
                          rtck=rtck, main_parttime=main_parttime)
            db.session.add(new_kit)
            db.session.commit()

            message = "Дані успішно додано!"

    return render_template('enter_data.html', message=message)


@app.route('/view_data', methods=['GET', 'POST'])
def view_data():
    if request.method == 'POST':
        last_name = request.form['last_name']

        results = Kit.query.filter(Kit.last_name.ilike(f"%{last_name}%")).all()

        if results:
            return render_template('view_data.html', results=results)
        else:
            return render_template('view_data.html', message="Записи не знайдено.")
    
    return render_template('view_data.html')

@app.route('/filter_data', methods=['GET', 'POST'])
def filter_data():
    kits = []  # Порожній список для результатів
    message = ""  # Повідомлення про помилки

    if request.method == 'POST':
        keywords = request.form.get('keywords', '').strip()
        age_input = request.form.get('age_range', '').strip()

        query = Kit.query  # Початковий запит без фільтрів

        # Фільтр за ключовими словами
        if keywords:
            query = query.filter(
                or_(
                    Kit.first_name.ilike(f'%{keywords}%'),
                    Kit.last_name.ilike(f'%{keywords}%'),
                    Kit.fathers_name.ilike(f'%{keywords}%'),
                    Kit.depart.ilike(f'%{keywords}%'),
                    Kit.position.ilike(f'%{keywords}%'),
                    Kit.rate.ilike(f'%{keywords}%'),
                    Kit.rtck.ilike(f'%{keywords}%'),
                    Kit.main_parttime.ilike(f'%{keywords}%')
                )
            )

        
        if age_input:
            try:
                if '-' in age_input:
                    min_age, max_age = map(int, age_input.split('-'))
                    query = query.filter(and_(Kit.age >= min_age, Kit.age <= max_age))
                else:
                    age_value = int(age_input)
                    query = query.filter(Kit.age == age_value)
            except ValueError:
                message = "Неправильний формат вікового введення! Введіть число (наприклад, 21) або діапазон (21-60)."

        kits = query.all()  
    
    return render_template('filter_data.html', kits=kits, message=message)

if __name__ == '__main__':
    app.run(debug=True)

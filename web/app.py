from flask import Flask,jsonify, render_template, request, redirect, url_for, send_file, session
import pandas as pd
import os
from datetime import datetime
from docx import Document
import zipfile
import io
app = Flask(__name__)
app.secret_key = os.urandom(24)
FILE_PATH = "data.xlsx"
TEMPLATE_PATH = "template.docx"
TEMPLATE_DIR = "templetes(1)"

def load_data():
    if not os.path.exists(FILE_PATH):
        df = pd.DataFrame(columns=['ID', 'First Name', 'Last Name', 'Fathers Name', 'Birth Date',
                                   'Age', 'Depart', 'Position', 'Rate', 'RTCK', 'Main/Part-time'])
        df.to_excel(FILE_PATH, index=False)
    return pd.read_excel(FILE_PATH)

def save_data(df):
    df.to_excel(FILE_PATH, index=False)

def calculate_age(birth_date):
    today = datetime.today().date()
    return today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/enter_data', methods=['GET', 'POST'])
def enter_data():
    message = ""
    df = load_data()
    if request.method == 'POST':
        id_input = request.form['id']
        first_name = request.form['first_name']
        last_name = request.form['last_name']
        fathers_name = request.form['fathers_name']
        birth_date = datetime.strptime(request.form['birth_date'], '%Y-%m-%d').date()
        age = calculate_age(birth_date)
        depart = request.form['depart']
        position = request.form['position']
        rate = float(request.form['rate'])
        rtck = request.form['rtck']
        main_parttime = request.form['main_parttime']
        if id_input in df['ID'].astype(str).values:
            existing_row = df[df['ID'].astype(str) == id_input].iloc[0]
            message = f"Дані з таким ID {id_input} вже існують: {existing_row['First Name']} {existing_row['Last Name']}, {existing_row['Position']}"
        else:
            new_row = pd.DataFrame([{ 'ID': id_input, 'First Name': first_name, 'Last Name': last_name, 'Fathers Name': fathers_name,
                'Birth Date': birth_date, 'Age': age, 'Depart': depart, 'Position': position,
                'Rate': rate, 'RTCK': rtck, 'Main/Part-time': main_parttime }])
            df = pd.concat([df, new_row], ignore_index=True)
            save_data(df)
            message = "Дані успішно додано!"
    return render_template('enter_data.html', message=message)


def fill_certificate(template_path, data, output_filename):
    """Заповнює шаблон документа переданими даними та зберігає його в пам'яті."""
    doc = Document(template_path)

    # Заміна тексту у всіх параграфах
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, data)

    # Заміна тексту у таблицях
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraph(cell.paragraphs[0], data)

    # Збереження документа в пам'яті
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)  # Потрібно повернутись на початок потоку
    return doc_stream

def replace_text_in_paragraph(paragraph, data):
    """Замінює всі маркери у переданому параграфі."""
    for key, value in data.items():
        placeholder = f"{{{key}}}"
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, str(value))

@app.route('/filter_data', methods=['GET', 'POST'])
def filter_data():
    df = load_data()
    filtered_df = df.copy()
    message = ""

    if request.method == 'POST':
        keywords = request.form.get('keywords', '').strip().lower()
        age_input = request.form.get('age_range', '').strip()

        if keywords:
            filtered_df = filtered_df[filtered_df.apply(lambda row: any(keywords in str(val).lower() for val in row), axis=1)]

        if age_input:
            try:
                if '-' in age_input:
                    min_age, max_age = map(int, age_input.split('-'))
                    filtered_df = filtered_df[(filtered_df['Age'] >= min_age) & (filtered_df['Age'] <= max_age)]
                else:
                    age_value = int(age_input)
                    filtered_df = filtered_df[filtered_df['Age'] == age_value]
            except ValueError:
                message = "Неправильний формат вікового введення!"

        session['filtered_data'] = filtered_df.to_dict(orient='records')

    templates = get_templates()
    return render_template('filter_data.html', kits=filtered_df.to_dict(orient='records'), message=message, templates=templates)

def get_templates():
    """Отримує список доступних шаблонів у папці."""
    if not os.path.exists(TEMPLATE_DIR):
        os.makedirs(TEMPLATE_DIR)
    return [f for f in os.listdir(TEMPLATE_DIR) if f.endswith(".docx")]


@app.route('/edit_data/<int:id>', methods=['GET', 'POST'])
def edit_data(id):
    df = load_data()
    record = df[df['ID'] == id].iloc[0]  # Get the record by ID

    if request.method == 'POST':
        # Get the updated data from the form
        first_name = request.form['first_name']
        last_name = request.form['last_name']
        fathers_name = request.form['fathers_name']
        birth_date = datetime.strptime(request.form['birth_date'], '%Y-%m-%d').date()
        age = calculate_age(birth_date)
        depart = request.form['depart']
        position = request.form['position']
        rate = float(request.form['rate'])
        rtck = request.form['rtck']
        main_parttime = request.form['main_parttime']

        # Update the record
        df.loc[df['ID'] == id, ['First Name', 'Last Name', 'Fathers Name', 'Birth Date', 'Age', 
                                 'Depart', 'Position', 'Rate', 'RTCK', 'Main/Part-time']] = [
            first_name, last_name, fathers_name, birth_date, age, depart, position, rate, rtck, main_parttime]
        save_data(df)

        return redirect(url_for('index'))  # Redirect to the main page after saving

    return render_template('edit_data.html', record=record)



@app.route('/generate_docs', methods=['POST'])
def generate_docs():
    """Генерує документи за оновленими даними"""
    filtered_df = pd.DataFrame(session.get('filtered_data', []))

    if filtered_df.empty:
        return "Немає даних для генерації документа.", 400

    selected_template = request.form.get("template")
    if not selected_template:
        return "Будь ласка, виберіть шаблон.", 400

    template_path = os.path.join(TEMPLATE_DIR, selected_template)

    # Створення архіву в пам'яті
    zip_stream = io.BytesIO()

    # Створення архіву
    with zipfile.ZipFile(zip_stream, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for _, row in filtered_df.iterrows():
            person_data = {
                "FirstName": row["First Name"],
                "LastName": row["Last Name"],
                "FathersName": row["Fathers Name"],
                "BirthDate": row["Birth Date"],
                "Age": row["Age"],
                "Depart": row["Depart"],
                "Position": row["Position"],
                "Rate": row["Rate"],
                "RTCK": row["RTCK"],
                "MainPartTime": row["Main/Part-time"],
            }
            filename = f"{row['Last Name']}_{row['First Name']}_certificate.docx"
            doc_stream = fill_certificate(template_path, person_data, filename)
            
            zipf.writestr(filename, doc_stream.read())  # Додаємо файл в архів

    zip_stream.seek(0)

    return send_file(zip_stream, as_attachment=True, download_name="Завантаження.zip", mimetype="application/zip")



if __name__ == '__main__':
    app.run(debug=True)
